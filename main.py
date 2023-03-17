import pdfplumber
import io
import pandas as pd
from docx import Document
from docx.shared import Pt
import pytesseract
from PIL import Image
import re
from tqdm import tqdm
import argparse

def extract_tables(page):
    tables = []
    for table in page.extract_tables():
        df = pd.DataFrame(table[1:], columns=table[0])
        tables.append(df)
    return tables

def create_table_in_docx(doc, df):
    table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
    for i, column in enumerate(df.columns):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(str(column) if column is not None else "")
        clear_formatting(run)

    for row_index, row in df.iterrows():
        for col_index, cell_value in enumerate(row):
            cell = table.cell(row_index+1, col_index)
            run = cell.paragraphs[0].add_run(str(cell_value) if cell_value is not None else "")
            clear_formatting(run)

    return table



def ocr_page(page):
    image = page.to_image().original
    ocr_text = pytesseract.image_to_string(image)
    return ocr_text


def clear_formatting(run, font_name='Arial', font_size=8):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = None
    run.font.italic = None
    run.font.underline = None
    run.font.strike = None
    run.font.superscript = None
    run.font.subscript = None

def is_sentence_end(text):
    return re.search(r"[.!?]\s*$", text)

def is_uppercase_first_char(text):
    return text[0].isupper()

def remove_empty_tables(doc):
    for table in doc.tables:
        is_empty = all(cell.text.strip() == "" for row in table.rows for cell in row.cells)
        if is_empty:
            remove_table(doc, table)

def remove_table(doc, table):
    tbl = table._element
    tbl.getparent().remove(tbl)
    tbl.clear()

def add_text_as_paragraphs(doc, text):
    paragraphs = text.split('\n')
    new_paragraphs = []

    current_paragraph = ""

    for paragraph_text in paragraphs:
        paragraph_text = paragraph_text.strip()

        if not paragraph_text:
            continue

        # Replace consecutive spaces with a single space
        paragraph_text = re.sub(r'\s{2,}', ' ', paragraph_text)
        # Replace "....." (continuous dots) with a tab
        paragraph_text = re.sub(r'\.{5,}', '\t', paragraph_text)
        # Replace continuous dots with spaces between them
        paragraph_text = re.sub(r'(\.(\s)*){2,}', '\t', paragraph_text)

        if not current_paragraph:
            current_paragraph = paragraph_text
        else:
            if is_sentence_end(current_paragraph) or is_uppercase_first_char(paragraph_text):
                new_paragraphs.append(current_paragraph)
                current_paragraph = paragraph_text
            else:
                current_paragraph += " " + paragraph_text

        # If the current paragraph ends with a sentence-ending punctuation mark or the next paragraph starts with an uppercase letter, add it to the new_paragraphs list
        if is_sentence_end(current_paragraph) or is_uppercase_first_char(paragraph_text):
            new_paragraphs.append(current_paragraph)
            current_paragraph = ""

    if current_paragraph:
        new_paragraphs.append(current_paragraph)

    for paragraph_text in new_paragraphs:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(paragraph_text)
        clear_formatting(run)
		
def pdf_to_docx(input_pdf, output_docx):
    pdf = pdfplumber.open(input_pdf)
    doc = Document()

    for page_num, page in enumerate(tqdm(pdf.pages, desc="Converting Pages")):
        # Add page number to the document
        if page_num > 0:
            doc.add_page_break()

        # Add page text
        text = page.extract_text()
        if not text:
            text = ocr_page(page)

        if text:
            add_text_as_paragraphs(doc, text)

        # Add tables
        tables = extract_tables(page)
        for df in tables:
            create_table_in_docx(doc, df)
            doc.add_paragraph()

    # Remove empty tables
    remove_empty_tables(doc)

    # Save the output DOCX file
    doc.save(output_docx)
    pdf.close()

def main(input_pdf, output_docx):
    pdf_to_docx(input_pdf, output_docx)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert PDF to DOCX")
    parser.add_argument("input_pdf", help="Input PDF file path")
    parser.add_argument("output_docx", help="Output DOCX file path")
    args = parser.parse_args()

    main(args.input_pdf, args.output_docx)